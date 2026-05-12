"""
Report Generator Module
=======================

生成完整的选品调研报告，支持多种输出格式。
"""

import json
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

from config import REPORT_CONFIG


class ReportGenerator:
    """报告生成器"""

    def __init__(self, data: dict, output_dir: Path):
        self.data = data
        self.output_dir = Path(output_dir)
        self.report_config = REPORT_CONFIG

        # 创建输出目录
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(self, format: str = "markdown") -> Path:
        """
        生成报告

        Args:
            format: 输出格式 (markdown, pdf, word, html)

        Returns:
            报告文件路径
        """
        # 生成文件名
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
        metadata = self.data.get("metadata", {})
        category = metadata.get("category", "report").replace(" ", "_")

        filename = f"{timestamp}_{category}_report"

        if format == "markdown":
            return self._generate_markdown(filename)
        elif format == "html":
            return self._generate_html(filename)
        elif format == "pdf":
            return self._generate_pdf(filename)
        elif format == "word":
            return self._generate_word(filename)
        else:
            return self._generate_markdown(filename)

    def _generate_markdown(self, filename: str) -> Path:
        """生成 Markdown 报告"""
        content = self._build_markdown_content()

        filepath = self.output_dir / f"{filename}.md"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)

        return filepath

    def _generate_html(self, filename: str) -> Path:
        """生成 HTML 报告"""
        md_content = self._build_markdown_content()

        # 简单的 Markdown to HTML 转换
        html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.data.get('metadata', {}).get('category', 'Product Research')} - 选品调研报告</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
               max-width: 900px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        h1 {{ color: #1a73e8; border-bottom: 2px solid #1a73e8; padding-bottom: 10px; }}
        h2 {{ color: #333; margin-top: 30px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 12px; text-align: left; }}
        th {{ background-color: #f5f5f5; }}
        .metric {{ display: inline-block; background: #e8f0fe; padding: 10px 20px;
                   border-radius: 8px; margin: 5px; }}
        .score {{ font-size: 24px; font-weight: bold; color: #1a73e8; }}
    </style>
</head>
<body>
    <div id="content"></div>
    <script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
    <script>
        document.getElementById('content').innerHTML = marked.parse(`{md_content}`);
    </script>
</body>
</html>"""

        filepath = self.output_dir / f"{filename}.html"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(html_content)

        return filepath

    def _generate_pdf(self, filename: str) -> Path:
        """生成 PDF 报告"""
        md_content = self._build_markdown_content()

        # 保存为 markdown 让用户自行转换
        md_path = self._generate_markdown(filename)

        print(f"   💡 PDF 转换提示: 使用 pandoc 转换")
        print(f"   pandoc {md_path} -o {filename}.pdf")

        return md_path

    def _generate_word(self, filename: str) -> Path:
        """生成 Word 报告"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            print("⚠️  请安装 python-docx: pip install python-docx")
            return self._generate_markdown(filename)

        doc = Document()

        # 标题
        title = doc.add_heading("亚马逊选品调研报告", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 元信息
        metadata = self.data.get("metadata", {})
        doc.add_paragraph(f"市场: {metadata.get('market', 'N/A')}")
        doc.add_paragraph(f"类目: {metadata.get('category', 'N/A')}")
        doc.add_paragraph(f"生成时间: {metadata.get('generated_at', 'N/A')}")

        # 添加各部分内容
        analysis = self.data.get("analysis", {})

        # 市场分析
        doc.add_heading("市场分析", level=1)
        market_analysis = analysis.get("market", {})
        if market_analysis:
            doc.add_paragraph(f"市场规模: {market_analysis.get('market_size', {}).get('estimated_size', 'N/A')}")
            doc.add_paragraph(f"品牌集中度: {market_analysis.get('brand_concentration', {}).get('top10_concentration', 'N/A')}")

        # 利润模型
        doc.add_heading("利润模型", level=1)
        profit_analysis = analysis.get("profit", {})
        if profit_analysis:
            tiers = profit_analysis.get("pricing_tiers", [])
            for tier in tiers:
                doc.add_paragraph(
                    f"{tier['name']}: ${tier['price']:.2f} "
                    f"(利润: ${tier['profit']:.2f}, 利润率: {tier['margin']*100:.1f}%)"
                )

        # TikTok 分析
        doc.add_heading("TikTok 传播分析", level=1)
        tiktok_analysis = analysis.get("tiktok", {})
        if tiktok_analysis:
            doc.add_paragraph(f"病毒传播评分: {tiktok_analysis.get('virality_score', 'N/A')}")

        filepath = self.output_dir / f"{filename}.docx"
        doc.save(str(filepath))

        return filepath

    def _build_markdown_content(self) -> str:
        """构建 Markdown 报告内容"""
        metadata = self.data.get("metadata", {})
        analysis = self.data.get("analysis", {})

        lines = []

        # 标题
        lines.append(f"# 🛒 亚马逊选品调研报告")
        lines.append("")
        lines.append(f"**市场:** {metadata.get('market', 'N/A')} | **类目:** {metadata.get('category', 'N/A')}")
        lines.append(f"**生成时间:** {metadata.get('generated_at', datetime.now().isoformat())}")
        lines.append("")
        lines.append("---")
        lines.append("")

        # 执行摘要
        lines.extend(self._build_executive_summary(analysis))

        # 市场分析
        lines.extend(self._build_market_section(analysis))

        # 竞争分析
        lines.extend(self._build_competition_section())

        # TikTok 验证
        lines.extend(self._build_tiktok_section(analysis))

        # 利润模型
        lines.extend(self._build_profit_section(analysis))

        # 选品建议
        lines.extend(self._build_recommendations_section(analysis))

        return "\n".join(lines)

    def _build_executive_summary(self, analysis: Dict) -> list:
        """构建执行摘要"""
        lines = []
        lines.append("## 📋 执行摘要")
        lines.append("")

        market_analysis = analysis.get("market", {})
        profit_analysis = analysis.get("profit", {})

        # 识别机会
        blue_ocean = market_analysis.get("blue_ocean_opportunities", [])
        if blue_ocean:
            lines.append("### 🎯 识别到的蓝海机会")
            lines.append("")
            for i, opp in enumerate(blue_ocean[:3], 1):
                lines.append(f"{i}. **{opp.get('title', 'Unknown')}**")
                lines.append(f"   - {opp.get('description', '')}")
                lines.append(f"   - 评分: {opp.get('score', 0)}/100")
                lines.append("")

        # 推荐定价
        if profit_analysis:
            recommended = profit_analysis.get("recommendations", {}).get("recommended_tier", {})
            if recommended:
                lines.append("### 💰 推荐定价方案")
                lines.append("")
                lines.append(f"| 产品定位 | 售价 | 利润 | 利润率 |")
                lines.append(f"|----------|------|------|--------|")
                lines.append(f"| {recommended.get('name', 'N/A')} | ${recommended.get('price', 0):.2f} | ${recommended.get('profit', 0):.2f} | {recommended.get('margin', 0)*100:.1f}% |")
                lines.append("")

        lines.append("---")
        lines.append("")

        return lines

    def _build_market_section(self, analysis: Dict) -> list:
        """构建市场分析章节"""
        lines = []
        lines.append("## 📊 市场分析")
        lines.append("")

        market_analysis = analysis.get("market", {})

        # 市场规模
        market_size = market_analysis.get("market_size", {})
        lines.append(f"### 市场规模")
        lines.append("")
        lines.append(f"- **市场规模:** {market_size.get('estimated_size', 'N/A')}")
        lines.append(f"- **年增长率 (CAGR):** {market_size.get('cagr', 0)*100:.1f}%")
        lines.append(f"- **商品总数:** {market_size.get('total_products', 0)}")
        lines.append("")

        # 品牌集中度
        brand_conc = market_analysis.get("brand_concentration", {})
        lines.append(f"### 品牌集中度")
        lines.append("")
        lines.append(f"- **独特品牌数:** {brand_conc.get('unique_brands', 0)}")
        lines.append(f"- **Top 10 集中度:** {brand_conc.get('top10_concentration', 0)*100:.1f}%")
        lines.append(f"- **建议:** {brand_conc.get('recommendation', 'N/A')}")
        lines.append("")

        lines.append("---")
        lines.append("")

        return lines

    def _build_competition_section(self) -> list:
        """构建竞争分析章节"""
        lines = []
        lines.append("## 🏆 竞争分析")
        lines.append("")
        lines.append("### 价格分布")
        lines.append("")

        # TODO: 接入真实数据
        lines.append("| 价格区间 | 商品数量 | 占比 |")
        lines.append("|----------|----------|------|")
        lines.append("| $0-10 | 20 | 20% |")
        lines.append("| $10-20 | 35 | 35% |")
        lines.append("| $20-30 | 25 | 25% |")
        lines.append("| $30+ | 20 | 20% |")
        lines.append("")

        lines.append("---")
        lines.append("")

        return lines

    def _build_tiktok_section(self, analysis: Dict) -> list:
        """构建 TikTok 验证章节"""
        lines = []
        lines.append("## 📱 TikTok 传播验证")
        lines.append("")

        tiktok_analysis = analysis.get("tiktok", {})

        # 病毒传播评分
        lines.append(f"### 病毒传播评分")
        lines.append("")
        score = tiktok_analysis.get("virality_score", 0)
        score_bar = "█" * int(score * 10) + "░" * (10 - int(score * 10))
        lines.append(f"`{score_bar}` {score*100:.0f}%")
        lines.append("")

        # 标签表现
        hashtags = tiktok_analysis.get("hashtag_performance", [])
        if hashtags:
            lines.append("### 标签表现")
            lines.append("")
            lines.append("| 标签 | 播放量 | 趋势 |")
            lines.append("|------|--------|------|")
            for tag in hashtags[:5]:
                lines.append(f"| {tag.get('hashtag', 'N/A')} | {tag.get('views', 0):,} | {tag.get('growth', 'N/A')} |")
            lines.append("")

        # 建议
        recommendations = tiktok_analysis.get("recommendations", [])
        if recommendations:
            lines.append("### 营销建议")
            lines.append("")
            for rec in recommendations:
                lines.append(f"- {rec}")
            lines.append("")

        lines.append("---")
        lines.append("")

        return lines

    def _build_profit_section(self, analysis: Dict) -> list:
        """构建利润模型章节"""
        lines = []
        lines.append("## 💰 利润模型")
        lines.append("")

        profit_analysis = analysis.get("profit", {})

        # 成本结构
        lines.append("### 成本结构 (以 $25 售价为例)")
        lines.append("")
        lines.append("| 成本项 | 金额 | 占比 |")
        lines.append("|--------|------|------|")
        lines.append("| 产品成本 | $5.80 | 23.2% |")
        lines.append("| 平台佣金 | $3.75 | 15.0% |")
        lines.append("| FBA 费用 | $3.50 | 14.0% |")
        lines.append("| 广告费 | $5.00 | 20.0% |")
        lines.append("| 退款损耗 | $0.75 | 3.0% |")
        lines.append("| 汇损 | $0.25 | 1.0% |")
        lines.append("| **总计** | **$19.05** | **76.2%** |")
        lines.append("")

        # 定价方案
        tiers = profit_analysis.get("pricing_tiers", [])
        if tiers:
            lines.append("### 定价方案")
            lines.append("")
            lines.append("| 产品定位 | 售价 | 成本 | 利润 | 毛利率 | 净利润率 |")
            lines.append("|----------|------|------|------|--------|----------|")
            for tier in tiers:
                lines.append(
                    f"| {tier.get('name', 'N/A')} | ${tier.get('price', 0):.2f} | "
                    f"${tier.get('cost', 0):.2f} | ${tier.get('profit', 0):.2f} | "
                    f"{(1 - tier.get('cost', 0)/tier.get('price', 1))*100:.1f}% | "
                    f"{tier.get('margin', 0)*100:.1f}% |"
                )
            lines.append("")

        lines.append("---")
        lines.append("")

        return lines

    def _build_recommendations_section(self, analysis: Dict) -> list:
        """构建选品建议章节"""
        lines = []
        lines.append("## ✅ 选品建议")
        lines.append("")
        lines.append("### 综合推荐")
        lines.append("")

        market_analysis = analysis.get("market", {})
        blue_ocean = market_analysis.get("blue_ocean_opportunities", [])

        if blue_ocean:
            for i, opp in enumerate(blue_ocean[:3], 1):
                lines.append(f"#### {i}. {opp.get('title', 'Unknown')}")
                lines.append("")
                lines.append(f"**描述:** {opp.get('description', '')}")
                lines.append("")
                lines.append(f"**评分:** {opp.get('score', 0)}/100")
                lines.append("")
        else:
            lines.append("基于当前数据分析，建议关注以下方向：")
            lines.append("")
            lines.append("1. **细分市场差异化** - 寻找头部品牌的弱点切入")
            lines.append("2. **内容营销驱动** - 利用 TikTok 等社交媒体建立品牌")
            lines.append("3. **品质驱动定价** - 不要陷入价格战，注重产品品质")
            lines.append("")

        # 风险提示
        lines.append("### ⚠️ 风险提示")
        lines.append("")
        lines.append("- 市场变化快，需持续关注竞品动态")
        lines.append("- 供应链稳定性至关重要，建议多供应商备选")
        lines.append("- 合规性检查，确保产品符合目标市场法规")
        lines.append("")

        lines.append("---")
        lines.append("")
        lines.append("*本报告由 Amazon Product Research Toolkit 自动生成*")
        lines.append(f"*生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")

        return lines
